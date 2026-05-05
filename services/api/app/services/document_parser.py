from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
import re
from typing import Any


TEXT_ENCODINGS = ("utf-8", "utf-8-sig", "gb18030", "latin-1")


class DocumentParserError(Exception):
    """Base class for deterministic document parser failures."""


class UnsupportedDocumentTypeError(DocumentParserError):
    """Raised when a document extension has no v0.1 basic parser."""


class EmptyParsedDocumentError(DocumentParserError):
    """Raised when a parser cannot extract any usable text."""


@dataclass(frozen=True)
class ParsedBlock:
    text: str
    content_type: str
    page_start: int | None = None
    page_end: int | None = None
    section_title: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(frozen=True)
class ParsedDocument:
    title: str | None
    language: str | None
    page_count: int | None
    blocks: list[ParsedBlock]
    metadata: dict[str, Any]


def normalize_extension(file_type: str | None, path: Path) -> str:
    extension = (file_type or path.suffix or "").lower()
    return extension if extension.startswith(".") else f".{extension}"


def normalize_text(text: str) -> str:
    return re.sub(r"\n{3,}", "\n\n", text.replace("\r\n", "\n").replace("\r", "\n")).strip()


def split_paragraphs(text: str) -> list[str]:
    normalized = normalize_text(text)
    if not normalized:
        return []
    return [part.strip() for part in re.split(r"\n\s*\n", normalized) if part.strip()]


def detect_language(text: str) -> str | None:
    compact_text = re.sub(r"\s+", "", text)
    if not compact_text:
        return None
    cjk_count = sum(1 for char in compact_text if "\u4e00" <= char <= "\u9fff")
    if cjk_count / max(len(compact_text), 1) > 0.2:
        return "zh-CN"
    ascii_letters = sum(1 for char in compact_text if ("a" <= char.lower() <= "z"))
    if ascii_letters >= 10:
        return "en"
    return None


def read_text_with_fallback(path: Path) -> tuple[str, str]:
    for encoding in TEXT_ENCODINGS:
        try:
            return path.read_text(encoding=encoding), encoding
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="latin-1"), "latin-1"


def finalize_parsed_document(
    *,
    title: str | None,
    page_count: int | None,
    blocks: list[ParsedBlock],
    metadata: dict[str, Any],
) -> ParsedDocument:
    clean_blocks = [
        ParsedBlock(
            text=normalize_text(block.text),
            content_type=block.content_type,
            page_start=block.page_start,
            page_end=block.page_end,
            section_title=block.section_title,
            metadata=block.metadata,
        )
        for block in blocks
        if normalize_text(block.text)
    ]
    if not clean_blocks:
        raise EmptyParsedDocumentError("No extractable text found.")
    full_text = "\n".join(block.text for block in clean_blocks)
    return ParsedDocument(
        title=title,
        language=detect_language(full_text),
        page_count=page_count,
        blocks=clean_blocks,
        metadata={
            **metadata,
            "block_count": len(clean_blocks),
            "text_char_count": len(full_text),
        },
    )


def parse_txt(path: Path) -> ParsedDocument:
    text, encoding = read_text_with_fallback(path)
    blocks = [ParsedBlock(text=paragraph, content_type="paragraph") for paragraph in split_paragraphs(text)]
    return finalize_parsed_document(
        title=None,
        page_count=None,
        blocks=blocks,
        metadata={"encoding": encoding, "parser": "txt_parser"},
    )


def parse_markdown(path: Path) -> ParsedDocument:
    text, encoding = read_text_with_fallback(path)
    blocks: list[ParsedBlock] = []
    paragraph_lines: list[str] = []
    current_heading: str | None = None
    title: str | None = None

    def flush_paragraph() -> None:
        if paragraph_lines:
            paragraph = normalize_text("\n".join(paragraph_lines))
            if paragraph:
                blocks.append(
                    ParsedBlock(
                        text=paragraph,
                        content_type="paragraph",
                        section_title=current_heading,
                    )
                )
            paragraph_lines.clear()

    for raw_line in text.splitlines():
        line = raw_line.strip()
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            flush_paragraph()
            current_heading = heading_match.group(2).strip()
            title = title or current_heading
            blocks.append(
                ParsedBlock(
                    text=current_heading,
                    content_type="heading",
                    section_title=current_heading,
                    metadata={"level": len(heading_match.group(1))},
                )
            )
        elif not line:
            flush_paragraph()
        else:
            paragraph_lines.append(line)
    flush_paragraph()

    return finalize_parsed_document(
        title=title,
        page_count=None,
        blocks=blocks,
        metadata={"encoding": encoding, "parser": "markdown_parser"},
    )


def parse_html(path: Path) -> ParsedDocument:
    from bs4 import BeautifulSoup

    text, encoding = read_text_with_fallback(path)
    soup = BeautifulSoup(text, "html.parser")
    blocks: list[ParsedBlock] = []
    title: str | None = None
    current_heading: str | None = None

    for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "table"]):
        content = normalize_text(element.get_text(" ", strip=True))
        if not content:
            continue
        if element.name and element.name.startswith("h"):
            current_heading = content
            title = title or content
            blocks.append(
                ParsedBlock(
                    text=content,
                    content_type="heading",
                    section_title=current_heading,
                    metadata={"tag": element.name},
                )
            )
        elif element.name == "li":
            blocks.append(ParsedBlock(text=content, content_type="list", section_title=current_heading))
        elif element.name == "table":
            blocks.append(ParsedBlock(text=content, content_type="table", section_title=current_heading))
        else:
            blocks.append(ParsedBlock(text=content, content_type="paragraph", section_title=current_heading))

    return finalize_parsed_document(
        title=title,
        page_count=None,
        blocks=blocks,
        metadata={"encoding": encoding, "parser": "html_parser"},
    )


def parse_docx(path: Path) -> ParsedDocument:
    from docx import Document

    document = Document(path)
    blocks: list[ParsedBlock] = []
    title: str | None = None
    current_heading: str | None = None

    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if "heading" in style_name.lower():
            current_heading = text
            title = title or text
            blocks.append(
                ParsedBlock(
                    text=text,
                    content_type="heading",
                    section_title=current_heading,
                    metadata={"style": style_name},
                )
            )
        else:
            blocks.append(ParsedBlock(text=text, content_type="paragraph", section_title=current_heading))

    for table_index, table in enumerate(document.tables):
        rows = []
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells if normalize_text(cell.text)]
            if cells:
                rows.append(" | ".join(cells))
        table_text = "\n".join(rows).strip()
        if table_text:
            blocks.append(
                ParsedBlock(
                    text=table_text,
                    content_type="table",
                    section_title=current_heading,
                    metadata={"table_index": table_index},
                )
            )

    return finalize_parsed_document(
        title=title,
        page_count=None,
        blocks=blocks,
        metadata={"parser": "docx_parser"},
    )


def parse_pdf(path: Path) -> ParsedDocument:
    from pypdf import PdfReader

    reader = PdfReader(path)
    blocks: list[ParsedBlock] = []
    for page_index, page in enumerate(reader.pages):
        text = normalize_text(page.extract_text() or "")
        if text:
            page_number = page_index + 1
            blocks.append(
                ParsedBlock(
                    text=text,
                    content_type="page",
                    page_start=page_number,
                    page_end=page_number,
                    metadata={"page_index": page_index},
                )
            )
    if not blocks:
        raise EmptyParsedDocumentError(
            "No extractable text found. OCR is not supported in v0.1 basic parser."
        )
    return finalize_parsed_document(
        title=None,
        page_count=len(reader.pages),
        blocks=blocks,
        metadata={"parser": "pdf_parser"},
    )


def parse_document_file(path: Path, file_type: str | None = None) -> ParsedDocument:
    extension = normalize_extension(file_type, path)
    if extension == ".txt":
        return parse_txt(path)
    if extension in {".md", ".markdown"}:
        return parse_markdown(path)
    if extension in {".html", ".htm"}:
        return parse_html(path)
    if extension == ".docx":
        return parse_docx(path)
    if extension == ".pdf":
        return parse_pdf(path)
    raise UnsupportedDocumentTypeError(f"Unsupported document parser extension: {extension or '<none>'}")
