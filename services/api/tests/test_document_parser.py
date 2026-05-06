from pathlib import Path

import pytest

from app.services.document_parser import (
    EmptyParsedDocumentError,
    UnsupportedDocumentTypeError,
    detect_language,
    parse_document_file,
)


def write_file(path: Path, content: str, encoding: str = "utf-8") -> Path:
    path.write_text(content, encoding=encoding)
    return path


def test_txt_parser_reads_text_and_detects_language(tmp_path) -> None:
    path = write_file(tmp_path / "policy.txt", "新能源政策正在推动市场变化。")

    parsed = parse_document_file(path, ".txt")

    assert parsed.language == "zh-CN"
    assert parsed.metadata["encoding"] == "utf-8"
    assert parsed.metadata["parser"] == "txt_parser"
    assert [block.text for block in parsed.blocks] == ["新能源政策正在推动市场变化。"]


def test_markdown_parser_preserves_headings(tmp_path) -> None:
    path = write_file(tmp_path / "memo.md", "# Main title\n\nIntro paragraph\n\n## Details\nMore text")

    parsed = parse_document_file(path, ".md")

    assert parsed.title == "Main title"
    assert [(block.content_type, block.text, block.section_title) for block in parsed.blocks] == [
        ("heading", "Main title", "Main title"),
        ("paragraph", "Intro paragraph", "Main title"),
        ("heading", "Details", "Details"),
        ("paragraph", "More text", "Details"),
    ]


def test_html_parser_extracts_headings_and_paragraphs(tmp_path) -> None:
    path = write_file(
        tmp_path / "memo.html",
        "<html><body><h1>Policy title</h1><p>First paragraph</p><ul><li>List item</li></ul></body></html>",
    )

    parsed = parse_document_file(path, ".html")

    assert parsed.title == "Policy title"
    assert [(block.content_type, block.text) for block in parsed.blocks] == [
        ("heading", "Policy title"),
        ("paragraph", "First paragraph"),
        ("list", "List item"),
    ]


def test_docx_parser_extracts_paragraphs(tmp_path) -> None:
    from docx import Document

    path = tmp_path / "memo.docx"
    doc = Document()
    doc.add_heading("Policy memo", level=1)
    doc.add_paragraph("DOCX paragraph")
    doc.save(path)

    parsed = parse_document_file(path, ".docx")

    assert parsed.title == "Policy memo"
    assert ("heading", "Policy memo") in [(block.content_type, block.text) for block in parsed.blocks]
    assert ("paragraph", "DOCX paragraph") in [(block.content_type, block.text) for block in parsed.blocks]


def test_pdf_parser_extracts_text_from_fixture(tmp_path) -> None:
    pytest.importorskip("pypdf")
    path = tmp_path / "text.pdf"
    path.write_bytes(
        b"%PDF-1.4\n"
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n"
        b"3 0 obj << /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> "
        b"/MediaBox [0 0 300 144] /Contents 5 0 R >> endobj\n"
        b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n"
        b"5 0 obj << /Length 44 >> stream\n"
        b"BT /F1 18 Tf 72 72 Td (Policy PDF text) Tj ET\n"
        b"endstream endobj\n"
        b"xref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n"
        b"0000000115 00000 n \n0000000257 00000 n \n0000000327 00000 n \n"
        b"trailer << /Root 1 0 R /Size 6 >>\nstartxref\n420\n%%EOF\n"
    )

    parsed = parse_document_file(path, ".pdf")

    assert parsed.page_count == 1
    assert parsed.blocks[0].content_type == "page"
    assert "Policy PDF text" in parsed.blocks[0].text


def test_unsupported_extension_raises(tmp_path) -> None:
    path = write_file(tmp_path / "memo.exe", "hello")

    with pytest.raises(UnsupportedDocumentTypeError):
        parse_document_file(path, ".exe")


def test_empty_text_raises(tmp_path) -> None:
    path = write_file(tmp_path / "empty.txt", "   \n")

    with pytest.raises(EmptyParsedDocumentError):
        parse_document_file(path, ".txt")


def test_detect_language_uses_simple_rules() -> None:
    assert detect_language("中国政策推动能源转型") == "zh-CN"
    assert detect_language("Policy research memo with market evidence") == "en"
    assert detect_language("12345") is None
